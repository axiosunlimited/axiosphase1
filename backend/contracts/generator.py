"""
Contract generation utilities using python-docx.
Generates employment contracts from template with filled-in values.
"""
from datetime import datetime
from io import BytesIO
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_contract_docx(data: Dict[str, Any]) -> BytesIO:
    """
    Generate an employment contract .docx file from template data.
    
    Args:
        data: Dictionary with contract fields:
            - employee_name (str)
            - national_id (str)
            - address (str)
            - mobile (str)
            - position (str)
            - department (str)
            - grade (str, optional)
            - contract_type (str)
            - start_date (str, YYYY-MM-DD)
            - end_date (str, YYYY-MM-DD)
            - probation_months (int, optional, default 3)
            - basic_salary (float)
            - transport_allowance (float, optional)
            - housing_allowance (float, optional)
            - bonus_enabled (bool, optional)
            - medical_aid_enabled (bool, optional)
            - school_fees_enabled (bool, optional)
            - reporting_to (str, optional)
            - signed_by (str, optional)
            - witness_name (str, optional)
    
    Returns:
        BytesIO: In-memory bytes of the generated .docx file
    """
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("ARRUPE JESUIT UNIVERSITY")
    header_run.font.size = Pt(14)
    header_run.font.bold = True
    
    doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CONTRACT OF EMPLOYMENT ENTERED INTO BY AND BETWEEN")
    title_run.font.bold = True
    
    doc.add_paragraph()
    
    # Employer info
    emp_title = doc.add_paragraph()
    emp_title_run = emp_title.add_run("ARRUPE JESUIT UNIVERSITY")
    emp_title_run.font.bold = True
    
    doc.add_paragraph("(Hereinafter referred to as \"the Employer\")")
    doc.add_paragraph("Located at 16 Link Road Mt Pleasant, Harare, Zimbabwe")
    doc.add_paragraph("Tel. +263 2424 745411")
    
    doc.add_paragraph()
    p = doc.add_paragraph("AND")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Employee info
    employee_name = data.get("employee_name", "").strip()
    doc.add_paragraph(f"{employee_name}")
    doc.add_paragraph("(Hereinafter referred to as \"the Employee\")")
    
    address = data.get("address", "").strip()
    doc.add_paragraph(f"Whose residential address is: {address}")
    
    national_id = data.get("national_id", "").strip()
    mobile = data.get("mobile", "").strip()
    doc.add_paragraph(f"National ID/Passport No: {national_id}   Mobile: {mobile}")
    
    doc.add_paragraph()
    
    # Background section
    doc.add_heading("Background", level=1)
    doc.add_paragraph("A. Your conditions of service are as set out in this document, which constitutes the contract of employment between you and Arrupe Jesuit University.")
    doc.add_paragraph("B. The terms and conditions of employment apply, other than as to matters provided for in the Labour Act Chapter 28.01 as amended from time to time.")
    
    contract_type = data.get("contract_type", "FIXED_TERM").upper()
    if contract_type == "PERMANENT":
        doc.add_paragraph("C. This is a permanent contract of employment subject to the terms and conditions set out herein.")
    else:
        doc.add_paragraph("C. This is a purely fixed term contract and shall not create legitimate expectation of renewal or permanent appointment.")
    
    doc.add_paragraph("D. AND WHEREAS the Employer and the Employee wish to reduce to writing the terms and conditions of their employment relationship...")
    
    doc.add_paragraph()
    
    # Section 1: Position and Department
    doc.add_heading("1. Position and Department", level=1)
    position = data.get("position", "").strip()
    department = data.get("department", "").strip()
    grade = data.get("grade", "").strip()
    
    grade_text = f" at {grade}" if grade else ""
    doc.add_paragraph(
        f"1.1 You shall be employed as a full time {position} in {department}{grade_text}. "
        "However, please note that the University reserves the right to transfer you to any other department in accordance with its needs."
    )
    
    doc.add_paragraph()
    
    # Section 2: Tenure
    doc.add_heading("2. Tenure", level=1)
    
    start_date = data.get("start_date", "")
    end_date = data.get("end_date", "")
    probation_months = data.get("probation_months", 3)
    
    # Format dates for readability
    try:
        start_obj = datetime.strptime(str(start_date), "%Y-%m-%d")
        start_formatted = start_obj.strftime("%d %B %Y")
    except:
        start_formatted = str(start_date)
    
    try:
        end_obj = datetime.strptime(str(end_date), "%Y-%m-%d")
        end_formatted = end_obj.strftime("%d %B %Y")
    except:
        end_formatted = str(end_date)
    
    if contract_type == "PERMANENT":
        doc.add_paragraph(
            f"2.1 Notwithstanding the date of writing and signing hereof, the employment contract shall be deemed to have come into effect from "
            f"the {start_formatted} for an indefinite period."
        )
    else:
        doc.add_paragraph(
            f"2.1 Notwithstanding the date of writing and signing hereof, the employment contract shall be deemed to have come into effect from "
            f"the {start_formatted} to the {end_formatted}. The contract is renewable subject to satisfactory performance."
        )
        
        doc.add_paragraph(
            "2.2 This is purely a fixed term contract which shall not create legitimate expectation of renewal or permanent appointment. "
            f"In this regard, the contract automatically lapses on {end_formatted}."
        )
    
    doc.add_paragraph(
        f"2.3 You shall be on {probation_months} months of probation, during which period the employer shall assess and evaluate your performance and conduct."
    )
    
    doc.add_paragraph("2.4 During the probation period, this contract may be terminated on one week notice by either party.")
    doc.add_paragraph("2.5 Upon completion of the probation period the employer shall explicitly communicate the results of the assessment and the decision to continue or not with the employment contract.")
    doc.add_paragraph("2.6 Upon confirmation after a successful probation, you shall continue with your responsibilities and duties as outlined in the detailed job description.")
    
    doc.add_paragraph("2.7 Upon confirmation after a successful probation, the appointment may be terminated:")
    doc.add_paragraph("2.7.1 by three months' notice given by yourself;", style='List Bullet')
    doc.add_paragraph("2.7.2 by mutual agreement;", style='List Bullet')
    doc.add_paragraph("2.7.3 for breach of an express or implied term of the contract upon such breach being verified after a due inquiry;", style='List Bullet')
    doc.add_paragraph("2.7.4 where applicable, on the recommendation of a Medical Board.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Section 3: Duties and Responsibilities
    doc.add_heading("3. Duties and Responsibilities", level=1)
    
    reporting_to = data.get("reporting_to", "the Dean of your Department").strip()
    doc.add_paragraph(f"3.1 You shall report directly to {reporting_to}.")
    doc.add_paragraph("3.2 Your responsibilities and duties are as outlined in the detailed job description attached to this contract.")
    doc.add_paragraph("3.3 By signing this contract, you agree that you shall at all times faithfully, industriously and to the best of your skills, ability, experience and talents perform all of the duties required of your position.")
    doc.add_paragraph("3.4 In carrying out these duties and responsibilities, you shall comply with all policies, procedures, rules and regulations of the Employer.")
    doc.add_paragraph("3.5 The University expects you to be exemplary, honest, respectful, accountable, trustworthy, committed, hardworking and loyal.")
    
    doc.add_paragraph()
    
    # Section 4: Hours of Work
    doc.add_heading("4. Hours of Work and Conflict of Interest", level=1)
    doc.add_paragraph("4.1 The hours of work are from 0800 hours to 1600 hours from Mondays to Fridays. Your lunch break shall be from 1230 hours to 1330 hours.")
    doc.add_paragraph("4.2 Your conduct with other staff members, students and associates of the University is expected to portray a high standard of professionalism and integrity.")
    doc.add_paragraph("4.3 Whilst in the employment of Arrupe Jesuit University, you may not engage in any other employment unless you have obtained the prior written consent of the University.")
    doc.add_paragraph("4.4 You shall not undertake any work which is likely to prejudice the interest of the University.")
    
    doc.add_paragraph()
    
    # Section 5: Remuneration
    doc.add_heading("5. Remuneration", level=1)
    
    salary = data.get("basic_salary", 0)
    try:
        salary_float = float(salary)
        salary_text = f"USD ${salary_float:,.2f}"
    except:
        salary_text = str(salary)
    
    doc.add_paragraph(f"5.1 Your basic salary shall be {salary_text} payable monthly in arrears.")
    doc.add_paragraph("5.2 Increments in salary may be made from time to time taking into account the University's financial capacity and prevailing economic situation.")
    
    doc.add_paragraph()
    
    # Section 6: Other Benefits
    doc.add_heading("6. Other Benefits", level=1)
    
    if data.get("bonus_enabled"):
        doc.add_paragraph("6.1 Bonus")
        doc.add_paragraph("6.1.1 Payment of bonus shall be at the discretion of the University and subject to availability of funds.", style='List Bullet')
        doc.add_paragraph("6.1.2 Subject to satisfactory performance, bonus (13th cheque) will be paid in November or December.", style='List Bullet')
        doc.add_paragraph()
    
    transport_allowance = data.get("transport_allowance", 120)
    if transport_allowance:
        doc.add_paragraph("6.2 Transport Allowance")
        doc.add_paragraph(f"6.2.1 You shall receive a monthly transport allowance of USD ${float(transport_allowance):.2f} which can be reviewed by management.")
        doc.add_paragraph()
    
    housing_allowance = data.get("housing_allowance", 175)
    if housing_allowance:
        doc.add_paragraph("6.3 Housing Allowance")
        doc.add_paragraph(f"6.3.1 You shall receive a housing allowance of USD ${float(housing_allowance):.2f} per month, if you are not residing in university premises.")
        doc.add_paragraph("6.3.2 Employees residing in any University houses or properties shall not receive a housing allowance.")
        doc.add_paragraph()
    
    if data.get("school_fees_enabled"):
        doc.add_paragraph("6.4 School Fees Allowance")
        doc.add_paragraph("6.4.1 You shall receive assistance for school fees for up to two direct/legal dependents at invoice value up to tertiary level.")
        doc.add_paragraph("6.4.2 Assistance shall be at 75% of specified school fees.")
        doc.add_paragraph()
    
    if data.get("medical_aid_enabled"):
        doc.add_paragraph("6.5 Medical Aid")
        doc.add_paragraph("6.5.1 You shall be required to register as a member of a Medical Aid Society through the University.")
        doc.add_paragraph("6.5.2 The University makes a 50% contribution for yourself, one spouse and three dependent children on the CIMAS Secure Package.")
        doc.add_paragraph()
    
    # Section 7: Leave
    doc.add_heading("7. Leave", level=1)
    doc.add_paragraph("7.1 Annual Leave")
    doc.add_paragraph("You shall be eligible for 2.5 days' vacation leave per month (30 days per year). Only up to 7 days may be carried forward to the next year.", style='List Number')
    doc.add_paragraph("7.2 Study Leave")
    doc.add_paragraph("You shall be entitled to at most 20 working days' study leave per year.", style='List Number')
    doc.add_paragraph("7.3 Sick Leave")
    doc.add_paragraph("You shall be entitled to up to 90 days sick leave on full pay and up to 90 days on half pay.", style='List Number')
    doc.add_paragraph("7.4 Special or Compassionate Leave")
    doc.add_paragraph("Special leave on full pay not exceeding 12 days in a calendar year shall be granted on justifiable grounds.", style='List Number')
    
    doc.add_paragraph()
    
    # Section 8: NSSA and Pension
    doc.add_heading("8. NSSA Contributions and RCCLE Pension Fund", level=1)
    doc.add_paragraph("8.1 You shall be required to contribute to the National Social Security Authority (NSSA) and RCCLE pension fund schemes as follows:")
    doc.add_paragraph("8.2 NSSA - 4.5% of basic monthly salary by yourself and 4.5% by the employer.", style='List Bullet')
    doc.add_paragraph("8.3 RCCLE Pension - 5% of basic monthly salary by yourself and 9% by the employer.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Section 9: Conditions Precedent
    doc.add_heading("9. Conditions Precedent", level=1)
    doc.add_paragraph("This offer is subject to:")
    doc.add_paragraph("Submission of proof of age and a police clearance report not later than two weeks upon taking up appointment;", style='List Bullet')
    doc.add_paragraph("Production of original academic certificates for verification no later than two days upon taking up appointment;", style='List Bullet')
    doc.add_paragraph("The requirements of the Immigration Act, if applicable;", style='List Bullet')
    doc.add_paragraph("Your dependants' birth certificates and other relevant documents.", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature section
    sig_title = doc.add_paragraph()
    sig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_run = sig_title.add_run("SIGNATURE PAGE")
    sig_run.font.bold = True
    sig_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Date line
    date_para = doc.add_paragraph("THUS, DONE AND SIGNED ON THIS DAY _____________________ 2024")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatory table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Row 1: Employer representative
    row1_cells = table.rows[0].cells
    row1_cells[0].text = "1. Employer\nRepresentative"
    row1_cells[1].text = "_____________________________"
    row1_cells[2].text = "Date: _____________"
    
    # Row 2: Witness
    row2_cells = table.rows[1].cells
    row2_cells[0].text = "2. Witness"
    witness_name = data.get("witness_name", "Human Resources Officer").strip()
    row2_cells[1].text = f"{witness_name}\n_____________________________"
    row2_cells[2].text = "Date: _____________"
    
    # Row 3: Employee
    row3_cells = table.rows[2].cells
    row3_cells[0].text = "3. Employee"
    row3_cells[1].text = f"{employee_name}\n_____________________________"
    row3_cells[2].text = "Date: _____________"
    
    # Row 4: Notes
    row4_cells = table.rows[3].cells
    row4_cells[0].text = "NB:"
    row4_cells[1].merge(row4_cells[2])
    row4_cells[1].text = "All signatures must be obtained in the presence of the other parties"
    
    # Return as BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
