from __future__ import annotations

import csv
from io import BytesIO, TextIOWrapper
from typing import Any, Dict, List, Tuple

import openpyxl

try:
    from docx import Document
except Exception:
    Document = None


ALLOWED_EXTS = {".csv", ".xlsx", ".docx"}


def _ext(name: str) -> str:
    name = (name or "").lower().strip()
    for e in ALLOWED_EXTS:
        if name.endswith(e):
            return e
    return ""


def parse_rows(file_obj) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Parse CSV/XLSX/DOCX (table) into a list of dict rows."""
    name = getattr(file_obj, "name", "")
    ext = _ext(name)
    if not ext:
        raise ValueError("Unsupported file type. Use CSV, XLSX, or DOCX.")

    if ext == ".csv":
        file_obj.seek(0)
        text = TextIOWrapper(file_obj, encoding="utf-8", errors="ignore")
        reader = csv.DictReader(text)
        headers = reader.fieldnames or []
        rows = [dict(r) for r in reader]
        return headers, rows

    if ext == ".xlsx":
        file_obj.seek(0)
        wb = openpyxl.load_workbook(file_obj, data_only=True)
        ws = wb.active
        headers = []
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(x).strip() if x is not None else "" for x in row]
                continue
            if all(x is None or str(x).strip() == "" for x in row):
                continue
            d = {}
            for j, h in enumerate(headers):
                if not h:
                    continue
                d[h] = row[j]
            rows.append(d)
        return headers, rows

    if ext == ".docx":
        if Document is None:
            raise ValueError("DOCX import requires python-docx.")
        file_obj.seek(0)
        doc = Document(file_obj)
        if not doc.tables:
            raise ValueError("DOCX must contain a structured table.")
        t = doc.tables[0]
        if len(t.rows) < 2:
            raise ValueError("DOCX table must include header + at least one data row.")
        headers = [c.text.strip() for c in t.rows[0].cells]
        rows = []
        for r in t.rows[1:]:
            vals = [c.text.strip() for c in r.cells]
            if all(v == "" for v in vals):
                continue
            d = {}
            for h, v in zip(headers, vals):
                if h:
                    d[h] = v
            rows.append(d)
        return headers, rows

    raise ValueError("Unsupported file type.")


def apply_mapping(row: Dict[str, Any], mapping: Dict[str, str]) -> Dict[str, Any]:
    """Map from file columns to target fields.

    mapping: {target_field: source_column}
    """
    if not mapping:
        # if no mapping provided, passthrough
        return dict(row)

    out: Dict[str, Any] = {}
    for target, source in mapping.items():
        if not source:
            continue
        out[target] = row.get(source)
    return out


def normalize_str(v: Any) -> str:
    if v is None:
        return ""
    return str(v).strip()
